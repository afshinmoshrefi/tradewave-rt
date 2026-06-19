"""Build the shareable agreement .docx from the canonical .md draft.

Usage: /home/flask/tradewave_realtime/.venv/bin/python make_agreement_docx.py
Pipeline: pandoc (faithful text conversion, smart punctuation OFF so no dash
or quote substitution ever touches legal text) -> python-docx styling pass
(fonts, margins, black headings, DRAFT footer with page numbers).
The .md stays the source of truth; rerun this after any edit to it.
"""
import re
import subprocess
import tempfile
from datetime import datetime
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SRC = "/home/flask/TradeWave_RT_AnneMarie_Agreement_DRAFT.md"
OUT = "/home/flask/TradeWave_RT_AnneMarie_Agreement_DRAFT.docx"
STAMP = (f"DRAFT for discussion - generated "
         f"{datetime.now(ZoneInfo('America/New_York')):%B %-d, %Y, %-I:%M %p ET}"
         f" - not legal advice, not for signature")

# Consecutive "1.1 ... / 1.2 ..." lines are one markdown paragraph; a contract
# needs each clause as its own paragraph. Blank-line them in a temp copy only.
text = open(SRC).read()
text = re.sub(r"\n(?=\d{1,2}\.\d+ )", "\n\n", text)
with tempfile.NamedTemporaryFile("w", suffix=".md", delete=False) as tf:
    tf.write(text)
    tmp_md = tf.name

subprocess.run(
    ["pandoc", "-f", "markdown-smart", "-t", "docx", tmp_md, "-o", OUT],
    check=True)

doc = Document(OUT)

# ---- page + base typography ----
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(1)
    s.left_margin = s.right_margin = Inches(1)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

# pandoc's style names trip python-docx's builtin-name translation on direct
# lookup; resolve by iterating instead.
by_name = {s.name: s for s in doc.styles}
HEADINGS = {"Heading 1": Pt(15), "Heading 2": Pt(13), "Heading 3": Pt(11.5)}
for name, size in HEADINGS.items():
    st = by_name[name]
    st.font.name = "Calibri"
    st.font.size = size
    st.font.bold = True
    st.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    st.paragraph_format.space_before = Pt(14)
    st.paragraph_format.space_after = Pt(6)

# the two read-first warning blocks (pandoc Block Text) read as callouts
if "Block Text" in by_name:
    by_name["Block Text"].font.size = Pt(10)
    by_name["Block Text"].font.italic = True

# ---- title block ----
title = doc.paragraphs[0]
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in title.runs:
    r.font.size = Pt(16)
sub = title.insert_paragraph_before("")  # placeholder, then move below title
# simpler: append a centered stamp line right after the title
sub._p.getparent().remove(sub._p)
stamp_p = doc.paragraphs[0]._p
new_p = OxmlElement("w:p")
stamp_p.addnext(new_p)
from docx.text.paragraph import Paragraph
sp = Paragraph(new_p, doc.paragraphs[0]._parent)
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sp.add_run(STAMP)
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(0x8B, 0x1A, 0x1A)

# ---- footer: stamp + Page X of Y ----
def field(p, instr):
    f = OxmlElement("w:fldSimple")
    f.set(qn("w:instr"), instr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    f.append(r)
    p._p.append(f)

for s in doc.sections:
    fp = s.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(STAMP + "  |  Page ")
    fr.font.size = Pt(8)
    field(fp, "PAGE")
    fr2 = fp.add_run(" of ")
    fr2.font.size = Pt(8)
    field(fp, "NUMPAGES")
    for r in fp.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ---- term-sheet table: tidy widths ----
for tbl in doc.tables:
    tbl.autofit = True
    if len(tbl.columns) == 2:
        for row in tbl.rows:
            row.cells[0].width = Inches(1.7)
            row.cells[1].width = Inches(4.8)

doc.save(OUT)
print("wrote", OUT)
